"""
tests/test_report_generator.py
Unit tests for DOCX report generation — no API key required.
"""

import sys
import os
import io
import unittest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from report_generator import generate_report

SAMPLE_REVIEW_RESULT = {
    "manuscript_title": "Prevalence of Osteoporosis in Elderly Agricultural Workers",
    "filename": "njcm-review-assignment-6266.docx",
    "word_count": 3412,
    "decision": "Major revision",
    "journal_name": "NJCM",
    "review_text": """\
---
PEER REVIEW REPORT
Manuscript Title: Prevalence of Osteoporosis in Elderly Agricultural Workers
Manuscript Type: Cross-sectional observational study
Date of Review: 2026-03-05
Reviewer: AI-Assisted Editorial Review System
---

STAGE 1: INITIAL EDITORIAL SCREENING
The manuscript includes title, abstract, keywords, and main sections.
Missing: ORCID IDs, CRediT taxonomy.

STAGE 2: SCOPE AND NOVELTY CHECK
Scope Fit: Strong
Genuine contribution to community medicine literature.

STAGE 3: METHODOLOGY REVIEW
MAJOR: Sample size calculation not documented.
MINOR: Blinding not applicable but should be stated.
SUGGESTION: Consider reporting confidence intervals for prevalence estimates.

STAGE 4: RESULTS AND DATA INTEGRITY
Results are clearly tabulated. Minor inconsistency in Table 2 row totals.

STAGE 5: DISCUSSION AND CONCLUSIONS
Conclusions align with results. Limitations section present but brief.

STAGE 6: REFERENCES
2 factual claims in Introduction lack citations.
Reference list formatting is inconsistent.

STAGE 7: ETHICAL AND INTEGRITY CHECKS
Ethics approval from institutional committee is stated. No integrity concerns.

STAGE 8: OVERALL EDITORIAL RECOMMENDATION
Decision: Major revision
Summary: The study makes a relevant contribution to community medicine.
Key corrections are required before acceptance.
Key Required Revisions:
1. Provide sample size calculation
2. Add ORCID IDs for all authors
3. Fix numerical inconsistency in Table 2
4. Add missing citations in Introduction

---
END OF REVIEW REPORT
---
""",
}


class TestReportGenerator(unittest.TestCase):

    def setUp(self):
        self.docx_bytes = generate_report(SAMPLE_REVIEW_RESULT)

    def test_generates_bytes(self):
        self.assertIsInstance(self.docx_bytes, bytes)

    def test_output_is_nonempty(self):
        self.assertGreater(len(self.docx_bytes), 1000)

    def test_output_is_valid_zip(self):
        """DOCX is a ZIP archive — must start with PK magic bytes."""
        self.assertEqual(self.docx_bytes[:2], b"PK")

    def test_output_is_openable_as_docx(self):
        """python-docx must be able to parse the generated file."""
        from docx import Document
        doc = Document(io.BytesIO(self.docx_bytes))
        self.assertIsNotNone(doc)

    def test_docx_contains_manuscript_title(self):
        from docx import Document
        doc = Document(io.BytesIO(self.docx_bytes))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        all_text = para_text + "\n" + table_text
        self.assertIn("Osteoporosis", all_text)

    def test_docx_contains_decision(self):
        from docx import Document
        doc = Document(io.BytesIO(self.docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Major revision", all_text)

    def test_accept_decision_generates_report(self):
        result = dict(SAMPLE_REVIEW_RESULT)
        result["decision"] = "Accept as is"
        result["review_text"] = result["review_text"].replace("Major revision", "Accept as is")
        docx_bytes = generate_report(result)
        self.assertGreater(len(docx_bytes), 1000)

    def test_reject_decision_generates_report(self):
        result = dict(SAMPLE_REVIEW_RESULT)
        result["decision"] = "Reject"
        docx_bytes = generate_report(result)
        self.assertGreater(len(docx_bytes), 1000)

    def test_missing_optional_fields_do_not_crash(self):
        minimal = {
            "manuscript_title": "Minimal Test",
            "filename": "test.docx",
            "word_count": 100,
            "decision": "Minor revision",
            "review_text": "STAGE 1: Initial screening complete.\nSTAGE 8: Decision: Minor revision",
        }
        docx_bytes = generate_report(minimal)
        self.assertGreater(len(docx_bytes), 500)


if __name__ == "__main__":
    unittest.main(verbosity=2)
