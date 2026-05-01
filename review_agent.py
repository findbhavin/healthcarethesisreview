"""
review_agent.py
Core AI review logic — extracts manuscript text, builds the system prompt
from guidelines/review_guidelines.yaml, and calls the Claude API.

The review framework (stages, checks, decision options) is entirely defined
in review_guidelines.yaml. Editors and SMEs can update that file without
touching this code.
"""

import os
import re
import logging
import anthropic
from guidelines.guidelines_loader import build_system_prompt, get_stage_weights, get_guidelines_version

logger = logging.getLogger(__name__)

_DEFAULT_MODELS = {
    "gemini": "gemini-2.5-pro",
    "anthropic": "claude-sonnet-4-6",
}


def _normalize_ai_config(ai_config: dict | None = None) -> dict:
    config = ai_config or {}
    provider = (config.get("provider") or os.environ.get("AI_PROVIDER") or "gemini").strip().lower()
    if provider not in ("gemini", "anthropic"):
        raise RuntimeError(f"Unsupported AI provider '{provider}'. Use 'gemini' or 'anthropic'.")

    model = (config.get("model") or os.environ.get("AI_MODEL") or _DEFAULT_MODELS[provider]).strip()
    if not model:
        model = _DEFAULT_MODELS[provider]

    anthropic_api_key = (config.get("anthropic_api_key") or os.environ.get("ANTHROPIC_API_KEY") or "").strip()
    gemini_api_key = (config.get("gemini_api_key") or os.environ.get("GEMINI_API_KEY") or "").strip()

    if provider == "anthropic" and not anthropic_api_key:
        raise RuntimeError("ANTHROPIC_API_KEY is not configured for Anthropic provider.")
    if provider == "gemini" and not gemini_api_key:
        raise RuntimeError("GEMINI_API_KEY is not configured for Gemini provider.")

    return {
        "provider": provider,
        "model": model,
        "anthropic_api_key": anthropic_api_key,
        "gemini_api_key": gemini_api_key,
    }


def generate_text(system_prompt: str, user_message: str, ai_config: dict | None = None) -> str:
    cfg = _normalize_ai_config(ai_config)
    provider = cfg["provider"]
    model = cfg["model"]

    if provider == "anthropic":
        client = anthropic.Anthropic(api_key=cfg["anthropic_api_key"])
        message = client.messages.create(
            model=model,
            max_tokens=8192,
            system=system_prompt,
            messages=[{"role": "user", "content": user_message}],
        )
        return message.content[0].text

    try:
        from google import genai as google_genai
        from google.genai import types as google_types
    except Exception as exc:
        raise RuntimeError(
            "Gemini SDK unavailable. Install 'google-genai' to use Gemini models."
        ) from exc

    client = google_genai.Client(api_key=cfg["gemini_api_key"])
    response = client.models.generate_content(
        model=model,
        contents=user_message,
        config=google_types.GenerateContentConfig(
            system_instruction=system_prompt,
            max_output_tokens=8192,
        ),
    )
    if not getattr(response, "text", None):
        raise RuntimeError("Gemini returned an empty response.")
    return response.text


# ---------------------------------------------------------------------------
# Text Extraction
# ---------------------------------------------------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file (paragraphs + tables)."""
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF file using pdfplumber."""
    import io
    import pdfplumber

    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)
    return "\n\n".join(parts)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to the appropriate extractor based on file extension."""
    lower = filename.lower()
    if lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(
            f"Unsupported file type: '{filename}'. "
            "Please upload a DOCX, PDF, or TXT file."
        )


# Cap very long manuscripts to keep token usage and latency reasonable.
_MAX_MANUSCRIPT_CHARS = 80_000


def _truncate(text: str) -> str:
    if len(text) <= _MAX_MANUSCRIPT_CHARS:
        return text
    logger.warning(
        f"Manuscript is {len(text):,} chars — truncating to {_MAX_MANUSCRIPT_CHARS:,} "
        "to reduce latency. Consider splitting very long documents."
    )
    return text[:_MAX_MANUSCRIPT_CHARS] + "\n\n[NOTE: Manuscript truncated at 80,000 characters for processing.]"


# ---------------------------------------------------------------------------
# Score Extraction
# ---------------------------------------------------------------------------

def _extract_stage_scores(review_text: str) -> dict:
    """
    Parse per-stage scores from the review text.
    Looks for patterns like:
      Score: 7/10  or  Score: 8/10  (within each STAGE block)

    Returns a dict:
      {
        "stage_scores": {1: 7, 2: 9, 3: 6, 4: 8, 5: 7, 6: 9, 7: 10},
        "weighted_score": 72.5,
        "wrs_display": "7×8 + 9×12 + ..."
      }
    """
    weights = get_stage_weights()
    stage_scores: dict[int, int] = {}

    # Find each STAGE N block and look for Score: X/10 within it
    # Use \d+ to handle two-digit stage numbers (10, 11)
    stage_pattern = re.compile(
        r"STAGE\s+(\d+)\s*[:\-—].*?(?=STAGE\s+\d+\s*[:\-—]|PRE-SUBMISSION REVIEW REPORT|END OF REVIEW|$)",
        re.DOTALL | re.IGNORECASE,
    )
    score_pattern = re.compile(r"Score\s*:\s*(\d+)\s*/\s*10", re.IGNORECASE)

    for match in stage_pattern.finditer(review_text):
        stage_num = int(match.group(1))
        block_text = match.group(0)
        score_match = score_pattern.search(block_text)
        if score_match:
            stage_scores[stage_num] = min(10, max(0, int(score_match.group(1))))

    # Also try to extract from a "WEIGHTED REVIEW SCORE" line in the report header
    wrs_line_match = re.search(
        r"WEIGHTED REVIEW SCORE.*?:\s*([\d.]+)\s*/\s*100",
        review_text, re.IGNORECASE,
    )

    # Compute weighted score from parsed stage scores
    # Stages 1–10 are scored; Stage 11 is the Final Recommendation (derived)
    weighted_sum = 0.0
    max_possible = 0.0
    wrs_parts = []
    for stage_num in range(1, 11):
        weight = weights.get(stage_num, 0)
        score = stage_scores.get(stage_num)
        if score is not None and weight > 0:
            weighted_sum += score * weight
            max_possible += 10 * weight
            wrs_parts.append(f"S{stage_num}={score}×{weight}")

    if max_possible > 0:
        weighted_score = round(weighted_sum / 10, 1)
    else:
        # Fall back to the parsed WRS line if available
        weighted_score = float(wrs_line_match.group(1)) if wrs_line_match else None

    return {
        "stage_scores": stage_scores,
        "weighted_score": weighted_score,
        "wrs_parts": " + ".join(wrs_parts),
    }


# ---------------------------------------------------------------------------
# Review Runner
# ---------------------------------------------------------------------------

def run_review(file_bytes: bytes, filename: str, journal_name: str = "",
               article_type: str = "", journal_tier: str = "",
               ai_config: dict | None = None) -> dict:
    """
    Run the AI peer review on an uploaded manuscript file.

    Parameters
    ----------
    file_bytes   : Raw bytes of the uploaded file
    filename     : Original filename (used to detect format)
    journal_name : Optional target journal name (used to inject journal-specific
                   requirements from review_guidelines.yaml)

    Returns
    -------
    dict with keys:
      manuscript_title : str
      review_text      : str   (full structured report)
      word_count       : int   (approximate)
      decision         : str   (parsed from report)
      filename         : str
    """
    manuscript_text = extract_text(file_bytes, filename)
    if not manuscript_text.strip():
        raise ValueError(
            "Could not extract any readable text from the uploaded file. "
            "Please ensure the file is not encrypted or image-only."
        )

    manuscript_text = _truncate(manuscript_text)
    word_count = len(manuscript_text.split())
    logger.info(
        f"Extracted {word_count} words from '{filename}' "
        f"(journal='{journal_name or 'not specified'}')"
    )

    # Build system prompt from guidelines YAML (picks up any SME edits)
    system_prompt = build_system_prompt(journal_name=journal_name)

    # Build context header for article type and journal tier
    context_lines = []
    if article_type:
        context_lines.append(f"Article Type: {article_type}")
    if journal_tier:
        context_lines.append(f"Target Journal Tier: {journal_tier}")
    context_header = ("\n".join(context_lines) + "\n\n") if context_lines else ""

    user_message = (
        f"{context_header}"
        f"Please conduct a full peer review of the following manuscript:\n\n"
        f"--- MANUSCRIPT BEGIN ---\n{manuscript_text}\n--- MANUSCRIPT END ---"
    )

    cfg = _normalize_ai_config(ai_config)
    logger.info(
        "Sending manuscript to %s model '%s' for review...",
        cfg["provider"],
        cfg["model"],
    )
    review_text = generate_text(system_prompt, user_message, ai_config=cfg)
    logger.info("Review completed successfully.")

    # Parse decision from report — handles both "Decision:" and "[X] Minor Revision" formats
    decision = "See report"
    for line in review_text.splitlines():
        stripped = line.strip()
        if stripped.lower().startswith("decision:"):
            decision = stripped.split(":", 1)[1].strip()
            break
        # New format: "PROBABLE DECISION:" followed by "[X] Minor Revision" on same or next lines
        if "probable decision" in stripped.lower():
            # check if decision is on this line or next
            after = stripped.split(":", 1)[1].strip() if ":" in stripped else ""
            if after:
                decision = after
            break
    # If still "See report", scan for checked decision box: "[X] Minor/Major/Rejection"
    if decision == "See report":
        for line in review_text.splitlines():
            m = re.search(r"\[X\]\s*(Minor Revision|Major Revision|Rejection)", line, re.IGNORECASE)
            if m:
                decision = m.group(1)
                break

    # Parse manuscript title from report
    manuscript_title = filename
    for line in review_text.splitlines():
        if "manuscript title" in line.lower() and ":" in line:
            candidate = line.split(":", 1)[1].strip()
            if candidate and candidate not in ("Not provided", "[Not provided]", "—", "[insert full title here]"):
                manuscript_title = candidate
            break

    scores = _extract_stage_scores(review_text)

    return {
        "manuscript_title": manuscript_title,
        "review_text": review_text,
        "word_count": word_count,
        "decision": decision,
        "filename": filename,
        "journal_name": journal_name,
        "article_type": article_type,
        "journal_tier": journal_tier,
        "stage_scores": scores["stage_scores"],
        "weighted_score": scores["weighted_score"],
        "wrs_parts": scores["wrs_parts"],
        "guidelines_version": get_guidelines_version(),
    }


def stream_review(file_bytes: bytes, filename: str, journal_name: str = "",
                  article_type: str = "", journal_tier: str = "",
                  ai_config: dict | None = None):
    """
    Streaming version of run_review. Yields dicts for SSE:
      {"type": "chunk", "text": "..."}          — incremental text
      {"type": "done",  "result": {...}}         — final result dict
      {"type": "error", "error": "..."}          — on failure

    The caller (Flask route) is responsible for JSON-encoding each dict.
    """
    try:
        manuscript_text = extract_text(file_bytes, filename)
        if not manuscript_text.strip():
            yield {"type": "error", "error": "Could not extract text from the file. Ensure it is not encrypted or image-only."}
            return

        manuscript_text = _truncate(manuscript_text)
        word_count = len(manuscript_text.split())
        logger.info(
            f"[stream] Extracted {word_count} words from '{filename}' "
            f"(journal='{journal_name or 'not specified'}')"
        )

        system_prompt = build_system_prompt(journal_name=journal_name)
        context_lines = []
        if article_type:
            context_lines.append(f"Article Type: {article_type}")
        if journal_tier:
            context_lines.append(f"Target Journal Tier: {journal_tier}")
        context_header = ("\n".join(context_lines) + "\n\n") if context_lines else ""
        user_message = (
            f"{context_header}"
            f"Please conduct a full peer review of the following manuscript:\n\n"
            f"--- MANUSCRIPT BEGIN ---\n{manuscript_text}\n--- MANUSCRIPT END ---"
        )

        cfg = _normalize_ai_config(ai_config)
        full_text_parts = []

        if cfg["provider"] == "anthropic":
            client = anthropic.Anthropic(api_key=cfg["anthropic_api_key"])
            logger.info("[stream] Streaming review from Anthropic model '%s'...", cfg["model"])
            with client.messages.stream(
                model=cfg["model"],
                max_tokens=8192,
                system=system_prompt,
                messages=[{"role": "user", "content": user_message}],
            ) as stream:
                for text_chunk in stream.text_stream:
                    full_text_parts.append(text_chunk)
                    yield {"type": "chunk", "text": text_chunk}
        else:
            logger.info("[stream] Generating review from Gemini model '%s'...", cfg["model"])
            review_text = generate_text(system_prompt, user_message, ai_config=cfg)
            full_text_parts.append(review_text)
            yield {"type": "chunk", "text": review_text}

        review_text = "".join(full_text_parts)
        logger.info("[stream] Review stream completed.")

        # Parse decision — handles "Decision:" and "[X] Minor/Major/Rejection" formats
        decision = "See report"
        for line in review_text.splitlines():
            stripped = line.strip()
            if stripped.lower().startswith("decision:"):
                decision = stripped.split(":", 1)[1].strip()
                break
            if "probable decision" in stripped.lower():
                after = stripped.split(":", 1)[1].strip() if ":" in stripped else ""
                if after:
                    decision = after
                break
        if decision == "See report":
            for line in review_text.splitlines():
                m = re.search(r"\[X\]\s*(Minor Revision|Major Revision|Rejection)", line, re.IGNORECASE)
                if m:
                    decision = m.group(1)
                    break

        # Parse title
        manuscript_title = filename
        for line in review_text.splitlines():
            if "manuscript title" in line.lower() and ":" in line:
                candidate = line.split(":", 1)[1].strip()
                if candidate and candidate not in ("Not provided", "[Not provided]", "—", "[insert full title here]"):
                    manuscript_title = candidate
                break

        scores = _extract_stage_scores(review_text)

        result = {
            "manuscript_title": manuscript_title,
            "review_text": review_text,
            "word_count": word_count,
            "decision": decision,
            "filename": filename,
            "journal_name": journal_name,
            "article_type": article_type,
            "journal_tier": journal_tier,
            "stage_scores": scores["stage_scores"],
            "weighted_score": scores["weighted_score"],
            "wrs_parts": scores["wrs_parts"],
        }
        yield {"type": "done", "result": result}

    except Exception as exc:
        logger.exception("[stream] Error during streaming review")
        yield {"type": "error", "error": str(exc)}
