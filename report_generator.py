"""
report_generator.py
Generates a formatted DOCX peer review report from the AI review output.
"""

import io
import re
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    """Set background colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.runs[0] if para.runs else para.add_run(text)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x96)  # dark blue
    return para


def _add_labelled_row(table, label: str, value: str, bg: str = "F0F4FA"):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    _set_cell_bg(row.cells[0], "DCE8F5")
    _set_cell_bg(row.cells[1], bg)
    row.cells[0].paragraphs[0].runs[0].bold = True


def _decision_color(decision: str) -> RGBColor:
    d = decision.lower()
    if "accept" in d:
        return RGBColor(0x00, 0x70, 0x00)  # green
    elif "minor" in d:
        return RGBColor(0xB8, 0x6E, 0x00)  # amber
    elif "major" in d:
        return RGBColor(0xCC, 0x33, 0x00)  # orange-red
    elif "reject" in d:
        return RGBColor(0xAA, 0x00, 0x00)  # red
    return RGBColor(0x00, 0x00, 0x00)


STAGE_HEADERS = [
    "STAGE 1",
    "STAGE 2",
    "STAGE 3",
    "STAGE 4",
    "STAGE 5",
    "STAGE 6",
    "STAGE 7",
    "STAGE 8",
]

STAGE_TITLES = {
    "STAGE 1": "Initial Editorial Screening",
    "STAGE 2": "Scope and Novelty Check",
    "STAGE 3": "Methodology Review",
    "STAGE 4": "Results and Data Integrity",
    "STAGE 5": "Discussion and Conclusions",
    "STAGE 6": "References",
    "STAGE 7": "Ethical and Integrity Checks",
    "STAGE 8": "Overall Editorial Recommendation",
}


def _split_into_stages(review_text: str) -> dict:
    """Split the review text into a dict keyed by stage tag (e.g. 'STAGE 1')."""
    stages = {}
    current_key = "PREAMBLE"
    current_lines = []

    for line in review_text.splitlines():
        matched = None
        for sh in STAGE_HEADERS:
            if line.strip().upper().startswith(sh):
                matched = sh
                break
        if matched:
            stages[current_key] = "\n".join(current_lines).strip()
            current_key = matched
            current_lines = [line]
        else:
            current_lines.append(line)

    stages[current_key] = "\n".join(current_lines).strip()
    return stages


def generate_report(review_result: dict) -> bytes:
    """
    Generate a DOCX report from the review result dict produced by review_agent.run_review().

    Returns raw bytes of the DOCX file.
    """
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # --- Title block ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("PEER REVIEW REPORT")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x96)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("AI-Assisted Editorial Review System")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub_run.italic = True

    doc.add_paragraph()

    # --- Manuscript info table ---
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"
    info_table.columns[0].width = Inches(2.0)
    info_table.columns[1].width = Inches(4.0)

    _add_labelled_row(info_table, "Manuscript Title", review_result.get("manuscript_title", "—"))
    _add_labelled_row(info_table, "File Name", review_result.get("filename", "—"))
    _add_labelled_row(info_table, "Word Count (approx.)", str(review_result.get("word_count", "—")))
    _add_labelled_row(info_table, "Date of Review", date.today().strftime("%B %d, %Y"))
    _add_labelled_row(info_table, "Reviewer", "AI-Assisted Editorial Review System")

    # Decision row with colour
    decision = review_result.get("decision", "See report")
    decision_row = info_table.add_row()
    decision_row.cells[0].text = "Editorial Decision"
    decision_row.cells[1].text = decision
    _set_cell_bg(decision_row.cells[0], "DCE8F5")
    _set_cell_bg(decision_row.cells[1], "FFFFFF")
    decision_row.cells[0].paragraphs[0].runs[0].bold = True
    if decision_row.cells[1].paragraphs[0].runs:
        decision_row.cells[1].paragraphs[0].runs[0].bold = True
        decision_row.cells[1].paragraphs[0].runs[0].font.color.rgb = _decision_color(decision)

    doc.add_paragraph()

    # --- Stage sections ---
    stages = _split_into_stages(review_result.get("review_text", ""))

    for stage_key in STAGE_HEADERS:
        stage_content = stages.get(stage_key, "")
        if not stage_content:
            continue

        title = STAGE_TITLES.get(stage_key, stage_key)
        _add_heading(doc, f"{stage_key}: {title}", level=2)

        # Remove the stage header line itself from content before printing
        lines = stage_content.splitlines()
        # Skip the first line if it's just repeating the stage header
        body_lines = []
        skip_first = True
        for line in lines:
            if skip_first and line.strip().upper().startswith(stage_key):
                skip_first = False
                continue
            skip_first = False
            body_lines.append(line)

        body = "\n".join(body_lines).strip()

        for paragraph_text in body.split("\n"):
            stripped = paragraph_text.strip()
            if not stripped:
                continue

            # Detect severity labels
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            severity_color = None
            severity_pattern = re.search(
                r"\b(MAJOR|MINOR|SUGGESTION)\b", stripped, re.IGNORECASE
            )
            if severity_pattern:
                sev = severity_pattern.group(1).upper()
                if sev == "MAJOR":
                    severity_color = RGBColor(0xCC, 0x00, 0x00)
                elif sev == "MINOR":
                    severity_color = RGBColor(0xB8, 0x6E, 0x00)
                elif sev == "SUGGESTION":
                    severity_color = RGBColor(0x00, 0x70, 0x00)

            run = p.add_run(stripped)
            run.font.size = Pt(10.5)
            if severity_color:
                run.font.color.rgb = severity_color

        doc.add_paragraph()

    # --- Footer note ---
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "This report was generated by an AI-assisted peer review system. "
        "It is intended to support, not replace, human editorial judgment."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
